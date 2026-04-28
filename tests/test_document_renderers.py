from __future__ import annotations

from importlib.metadata import version as package_version
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import struct
import zlib
import zipfile

import docscriptor
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from pypdf import PdfReader

import docscriptor.components.generated as generated_components
import docscriptor.components.inline as inline_components
from docscriptor.layout.indexing import build_render_index
from docscriptor import (
    Affiliation,
    Author,
    AuthorLayout,
    Box,
    BoxStyle,
    BulletList,
    CitationLibrary,
    CitationSource,
    Chapter,
    Comment,
    CommentsPage,
    CodeBlock,
    Document,
    DocumentSettings,
    Equation,
    Figure,
    FigureList,
    Footnote,
    HeadingNumbering,
    ListStyle,
    Math,
    NumberedList,
    Paragraph,
    ParagraphStyle,
    ReferencesPage,
    Section,
    Subsection,
    Subsubsection,
    Table,
    TableCell,
    TableStyle,
    TableOfContents,
    TableList,
    Text,
    Theme,
    bold,
    code,
    color,
    cite,
    comment,
    footnote,
    italic,
    link,
    math,
    markup,
    styled,
)
from docscriptor.settings import TextStyle

class HighlightedParagraph(Paragraph):
    pass


def _write_sample_image(path: Path) -> None:
    path.write_bytes(_build_sample_png())


def _build_sample_png(width: int = 360, height: int = 220) -> bytes:
    rows: list[bytes] = []
    for y in range(height):
        row = bytearray([0])
        for x in range(width):
            if y < 34:
                pixel = (34, 58, 94)
            elif x < 18 or x >= width - 18 or y < 52 or y >= height - 18:
                pixel = (214, 221, 233)
            elif 26 < x < width - 26 and 70 < y < 102:
                pixel = (205, 121, 62)
            elif (x - 36) // 54 % 2 == 0 and 122 < y < 182:
                pixel = (89, 132, 198)
            else:
                pixel = (247, 249, 252)
            row.extend(pixel)
        rows.append(bytes(row))

    raw_image = b"".join(rows)
    return b"".join(
        (
            b"\x89PNG\r\n\x1a\n",
            _png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)),
            _png_chunk(b"IDAT", zlib.compress(raw_image, level=9)),
            _png_chunk(b"IEND", b""),
        )
    )


def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    payload = chunk_type + data
    checksum = zlib.crc32(payload) & 0xFFFFFFFF
    return struct.pack(">I", len(data)) + payload + struct.pack(">I", checksum)


class FakeAxis:
    def __init__(self, values: list[object], names: tuple[str | None, ...] = ("",)) -> None:
        self._values = values
        self.names = names
        self.name = names[0] if names else None
        self.nlevels = max((len(value) if isinstance(value, tuple) else 1) for value in values) if values else 1

    def tolist(self) -> list[object]:
        return list(self._values)

    def __iter__(self):
        return iter(self._values)


class FakeDataFrame:
    def __init__(
        self,
        *,
        columns: list[object],
        rows: list[list[object]],
        index: list[object] | None = None,
        index_names: tuple[str | None, ...] = ("",),
    ) -> None:
        self.columns = FakeAxis(columns)
        self._rows = rows
        self.index = FakeAxis(index or list(range(len(rows))), names=index_names)

    def itertuples(self, *, index: bool = False, name: str | None = None):
        for row_index, row in enumerate(self._rows):
            if index:
                yield (self.index.tolist()[row_index], *row)
            else:
                yield tuple(row)


class FakeFigure:
    def __init__(self, image_bytes: bytes) -> None:
        self.image_bytes = image_bytes
        self.calls: list[dict[str, object]] = []

    def savefig(self, target: object, **kwargs: object) -> None:
        self.calls.append(dict(kwargs))
        target.write(self.image_bytes)


def _pdf_font_names(pdf_path: Path) -> set[str]:
    font_names: set[str] = set()
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/Font" not in resources:
            continue
        fonts = resources["/Font"].get_object()
        for font in fonts.values():
            font_object = font.get_object()
            base_font = font_object.get("/BaseFont")
            if base_font is not None:
                font_names.add(str(base_font))
    return font_names


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _pdf_content_bytes(pdf_path: Path) -> bytes:
    parts: list[bytes] = []
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        contents = page.get_contents()
        if contents is None:
            continue
        if isinstance(contents, list):
            for item in contents:
                parts.append(item.get_data())
        else:
            parts.append(contents.get_data())
    return b"\n".join(parts)


def _pdf_text_context(pdf_path: Path, text: str, window: int = 160) -> bytes:
    content = _pdf_content_bytes(pdf_path)
    needle = f"({text})".encode()
    index = content.find(needle)
    assert index != -1, f"{text!r} not found in PDF content stream"
    start = max(index - window, 0)
    return content[start : index + len(needle) + window]


def _docx_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_version_is_defined() -> None:
    from docscriptor import __version__

    assert __version__ == package_version("docscriptor")
    assert __version__


def test_markup_creates_styled_fragments() -> None:
    fragments = markup("plain **bold** *italic* `code`")

    assert [fragment.value for fragment in fragments] == ["plain ", "bold", " ", "italic", " ", "code"]
    assert fragments[1].style.bold is True
    assert fragments[3].style.italic is True
    assert fragments[5].style.font_name == "Courier New"


def test_list_classes_create_block_instances() -> None:
    bullet = BulletList("first", Paragraph("second"))
    ordered = NumberedList("step one", "step two")

    assert isinstance(bullet, BulletList)
    assert bullet.ordered is False
    assert [item.plain_text() for item in bullet.items] == ["first", "second"]
    assert isinstance(ordered, NumberedList)
    assert ordered.ordered is True
    assert [item.plain_text() for item in ordered.items] == ["step one", "step two"]


def test_comment_and_math_helpers_create_renderable_fragments() -> None:
    inline_comment = comment("term", "Expanded note", author="pytest", initials="PT")
    inline_footnote = footnote("term", "Portable footnote note")
    inline_math = math(r"\alpha^2 + \beta^2")
    equation = Equation(r"\frac{1}{2}")

    assert isinstance(inline_comment, docscriptor.Comment)
    assert inline_comment.plain_text() == "term[?]"
    assert inline_comment.author == "pytest"
    assert inline_comment.initials == "PT"
    assert isinstance(inline_footnote, Footnote)
    assert inline_footnote.plain_text() == "term[?]"
    assert isinstance(inline_math, Math)
    assert inline_math.plain_text() == "alpha2 + beta2"
    assert equation.plain_text() == "(1)/(2)"


def test_method_style_inline_actions_create_renderable_fragments() -> None:
    source = CitationSource("Usage Guide", key="guide", year="2026")
    library = CitationLibrary([source])

    assert Text.bold("important").style.bold is True
    assert Text.italic("note").style.italic is True
    assert Text.code("x = 1").style.font_name == "Courier New"
    assert Text.color("accent", "#0066AA").style.color == "0066AA"
    assert [fragment.value for fragment in Text.from_markup("plain **bold**")] == [
        "plain ",
        "bold",
    ]
    assert bold("important").style.bold is True
    assert italic("note").style.italic is True
    assert code("x = 1").style.font_name == "Courier New"
    assert color("accent", "#0066AA").style.color == "0066AA"
    external_link = link("https://example.com", "Example")
    assert isinstance(external_link, inline_components.Hyperlink)
    assert external_link.target == "https://example.com"
    assert external_link.plain_text() == "Example"
    assert source.cite().plain_text() == "[?]"
    assert library.cite("guide").plain_text() == "[?]"
    assert Comment.annotated("term", "Expanded note").plain_text() == "term[?]"
    assert Footnote.annotated("term", "Portable footnote note").plain_text() == "term[?]"
    assert Math.inline(r"\alpha^2").plain_text() == "alpha2"


def test_theme_validates_page_number_configuration() -> None:
    theme = Theme(show_page_numbers=True, page_number_format="Page {page}", page_number_alignment="right")

    assert theme.format_page_number(3) == "Page 3"
    assert theme.format_page_number(3, front_matter=True) == "Page iii"

    try:
        Theme(page_number_format="Page")
    except ValueError as exc:
        assert "{page}" in str(exc)
    else:
        raise AssertionError("Expected page_number_format validation to fail")


def test_paragraph_style_defaults_to_justify_alignment() -> None:
    paragraph = Paragraph("Default alignment paragraph.")

    assert paragraph.style.alignment == "justify"


def test_theme_defaults_center_media_objects_and_captions() -> None:
    theme = Theme()

    assert theme.page_background_color == "FFFFFF"
    assert theme.caption_alignment == "center"
    assert theme.table_alignment == "center"
    assert theme.figure_alignment == "center"
    assert theme.box_alignment == "center"


def test_page_background_color_renders_to_all_outputs(tmp_path: Path) -> None:
    document = Document(
        "Color Test",
        Paragraph("Tinted page."),
        settings=DocumentSettings(theme=Theme(page_background_color="#F4F8FC")),
    )
    docx_path = tmp_path / "color.docx"
    pdf_path = tmp_path / "color.pdf"
    html_path = tmp_path / "color.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    assert 'w:background w:color="F4F8FC"' in _docx_document_xml(docx_path)
    assert b".956863 .972549 .988235 rg" in _pdf_content_bytes(pdf_path)
    assert "background: #F4F8FC;" in html_path.read_text(encoding="utf-8")


def test_numbering_and_list_styles_are_customizable() -> None:
    heading_numbering = HeadingNumbering(formats=("upper-roman", "lower-alpha"), prefix="[", suffix="]")
    ordered_style = ListStyle(marker_format="upper-roman", prefix="(", suffix=")")
    bullet_style = ListStyle(marker_format="bullet", bullet="•", suffix="")

    assert heading_numbering.format_label([2, 3]) == "[II.c]"
    assert ordered_style.marker_for(0) == "(I)"
    assert ordered_style.marker_for(2) == "(III)"
    assert bullet_style.marker_for(1) == "•"


def test_table_accepts_dataframe_like_inputs_and_spans() -> None:
    dataframe = FakeDataFrame(
        columns=[("Metrics", "Latency"), ("Metrics", "Quality"), ("Summary", "")],
        rows=[["14 ms", "stable", "ready"]],
    )
    table = Table(
        dataframe,
        caption="Span test.",
        column_widths=[1.5, 1.5, 1.5],
        style=TableStyle(alternate_row_background_color="#F4F8FC"),
    )
    merged_header = Table(
        headers=[
            [TableCell("Metrics", colspan=2), TableCell("Summary", rowspan=2)],
            ["Latency", "Quality"],
        ],
        rows=[["14 ms", "stable", "ready"]],
        column_widths=[1.5, 1.5, 1.5],
    )

    assert len(table.header_rows) == 2
    assert table.header_rows[0][0].colspan == 2
    assert table.header_rows[0][1].rowspan == 2
    assert table.layout().column_count == 3
    assert merged_header.layout().row_count == 3


def test_heading_hierarchy_uses_latex_like_levels() -> None:
    chapter = Chapter(
        "Part I",
        Section(
            "Overview",
            Subsection(
                "Details",
                Subsubsection("Examples"),
            ),
        ),
    )

    assert chapter.level == 1
    assert chapter.children[0].level == 2
    assert chapter.children[0].children[0].level == 3
    assert chapter.children[0].children[0].children[0].level == 4


def test_sections_can_be_rendered_without_numbering() -> None:
    unnumbered = Section("Abstract", Paragraph("Summary."), level=2, numbered=False)
    numbered = Section("Introduction", Paragraph("Body."), level=1)
    document = Document("Article", unnumbered, numbered)

    render_index = build_render_index(document)

    assert render_index.heading_number(unnumbered) is None
    assert render_index.heading_number(numbered) == "1"


def test_public_api_prefers_classes_for_structural_nodes() -> None:
    assert hasattr(docscriptor, "Document")
    assert hasattr(docscriptor, "DocumentSettings")
    assert hasattr(docscriptor, "Chapter")
    assert hasattr(docscriptor, "AuthorLayout")
    assert hasattr(docscriptor, "Section")
    assert hasattr(docscriptor, "Paragraph")
    assert hasattr(docscriptor, "BulletList")
    assert hasattr(docscriptor, "NumberedList")
    assert hasattr(docscriptor, "TableList")
    assert hasattr(docscriptor, "FigureList")
    assert hasattr(docscriptor, "cite")
    assert hasattr(docscriptor, "Box")
    assert hasattr(docscriptor, "BoxStyle")
    assert hasattr(docscriptor, "HeadingNumbering")
    assert hasattr(docscriptor, "ListStyle")
    assert hasattr(docscriptor, "Table")
    assert hasattr(docscriptor, "TableCell")
    assert hasattr(docscriptor, "TableStyle")
    assert hasattr(docscriptor, "Figure")
    assert hasattr(docscriptor, "TableOfContents")
    assert hasattr(docscriptor, "Comment")
    assert hasattr(docscriptor, "CommentsPage")
    assert hasattr(docscriptor, "Footnote")
    assert hasattr(docscriptor, "Equation")
    assert hasattr(docscriptor, "Math")
    assert hasattr(docscriptor, "TextStyle")
    assert not hasattr(docscriptor, "Body")
    assert not hasattr(docscriptor, "Bold")
    assert not hasattr(docscriptor, "Hyperlink")
    assert not hasattr(docscriptor, "Italic")
    assert not hasattr(docscriptor, "Monospace")
    assert not hasattr(docscriptor, "FootnotesPage")
    assert hasattr(inline_components, "Bold")
    assert hasattr(inline_components, "Hyperlink")
    assert hasattr(inline_components, "Italic")
    assert hasattr(inline_components, "Monospace")
    assert hasattr(generated_components, "FootnotesPage")
    assert not hasattr(docscriptor, "ListBlock")
    assert not hasattr(docscriptor, "Citation")
    assert not hasattr(docscriptor, "TableReference")
    assert not hasattr(docscriptor, "FigureReference")
    assert not hasattr(docscriptor, "Strong")
    assert not hasattr(docscriptor, "Emphasis")
    assert not hasattr(docscriptor, "Code")
    assert hasattr(docscriptor, "comment")
    assert hasattr(docscriptor, "footnote")
    assert hasattr(docscriptor, "math")
    assert hasattr(docscriptor, "bold")
    assert hasattr(docscriptor, "italic")
    assert hasattr(docscriptor, "code")
    assert hasattr(docscriptor, "color")
    assert hasattr(docscriptor, "link")

    for removed_name in (
        "document",
        "body",
        "chapter",
        "section",
        "subsection",
        "subsubsection",
        "paragraph",
        "code_block",
        "bullet_list",
        "numbered_list",
        "table",
        "figure",
    ):
        assert not hasattr(docscriptor, removed_name)


def test_bibtex_string_creates_citation_library() -> None:
    document = Document(
        "Bibliography Test",
        Paragraph("Example"),
        citations="""@misc{pydocs-repository,
  title = {pydocs},
  organization = {Gonie-Gonie},
  year = {2026},
  url = {https://github.com/Gonie-Gonie/pydocs},
  note = {GitHub repository}
}""",
    )

    entry = document.citations.resolve("pydocs-repository")
    assert entry.title == "pydocs"
    assert entry.organization == "Gonie-Gonie"
    assert entry.url == "https://github.com/Gonie-Gonie/pydocs"
    assert "GitHub repository" in entry.format_reference()


def test_document_accepts_document_settings() -> None:
    settings = DocumentSettings(
        author="Docscriptor",
        summary="Settings test",
        subtitle="Grouped metadata",
        authors=[
            Author(
                "Example Author",
                affiliations=[Affiliation(organization="Example Lab")],
            )
        ],
        author_layout=AuthorLayout(mode="stacked"),
        cover_page=True,
        theme=Theme(show_page_numbers=True),
    )

    document = Document("Configured", Paragraph("Body"), settings=settings)

    assert document.author == "Docscriptor"
    assert document.summary == "Settings test"
    assert document.subtitle is not None
    assert document.subtitle[0].plain_text() == "Grouped metadata"
    assert document.authors[0].name == "Example Author"
    assert document.authors[0].affiliations[0].formatted() == "Example Lab"
    assert document.settings.author_layout.mode == "stacked"
    assert document.cover_page is True
    assert document.theme.show_page_numbers is True


def test_auto_footnotes_page_can_be_disabled(tmp_path: Path) -> None:
    document = Document(
        "Inline Notes",
        Paragraph(
            "Portable ",
            footnote("term", "Suppressed footnote note."),
            " stays inline.",
        ),
        settings=DocumentSettings(
            theme=Theme(
                auto_footnotes_page=False,
                footnote_placement="document",
            )
        ),
    )

    docx_path = tmp_path / "inline-notes.docx"
    pdf_path = tmp_path / "inline-notes.pdf"
    html_path = tmp_path / "inline-notes.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    normalized_html_text = _normalized_html_text(html_path)

    assert "Footnotes" not in paragraph_texts
    assert all("Suppressed footnote note." not in text for text in paragraph_texts)
    assert "Footnotes" not in pdf_text
    assert "Suppressed footnote note." not in pdf_text
    assert "Footnotes" not in normalized_html_text
    assert "Suppressed footnote note." not in normalized_html_text


def test_docx_native_page_footnotes_are_default(tmp_path: Path) -> None:
    document = Document(
        "Inline Notes",
        Paragraph(
            "Portable ",
            footnote("term", "Default native footnote note."),
            " stays inline.",
        ),
    )

    docx_path = tmp_path / "inline-notes.docx"
    document.save_docx(docx_path)

    paragraph_texts = [paragraph.text for paragraph in WordDocument(docx_path).paragraphs]
    assert "Footnotes" not in paragraph_texts
    assert all("Default native footnote note." not in text for text in paragraph_texts)

    with zipfile.ZipFile(docx_path) as archive:
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
    assert "Default native footnote note." in footnotes_xml
    assert "w:footnoteReference" in _docx_document_xml(docx_path)


def test_document_renders_to_docx_and_pdf(tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    _write_sample_image(image_path)
    repository_source = CitationSource(
        "pydocs",
        organization="Gonie-Gonie",
        publisher="GitHub repository",
        year="2026",
        url="https://github.com/Gonie-Gonie/pydocs",
    )
    registered_source = CitationSource(
        "Release Notes",
        key="release-notes",
        organization="Gonie-Gonie",
        publisher="Documentation index",
        year="2026",
        url="https://github.com/Gonie-Gonie/pydocs/releases",
    )
    unused_source = CitationSource(
        "Internal Draft",
        key="internal-draft",
        organization="Gonie-Gonie",
        year="2026",
        url="https://example.invalid/internal-draft",
    )
    artifacts_table = Table(
        headers=["Type", "Status"],
        rows=[
            ["DOCX", Paragraph("generated ", footnote("state", "Table cell footnote note."))],
            ["PDF", "generated"],
        ],
        caption="Generated artifacts.",
        column_widths=[2.5, 2.5],
        style=TableStyle(
            header_background_color="#DCE8F4",
            alternate_row_background_color="#F7FAFD",
        ),
    )
    workflow_frame = FakeDataFrame(
        columns=[("Workflow", "Step"), ("Workflow", "Target"), ("Result", "")],
        rows=[
            ["Draft review", "DOCX", "ready"],
            ["Release", "PDF", "published"],
        ],
    )
    workflow_table = Table(
        workflow_frame,
        caption="Output workflow.",
        column_widths=[2.2, 1.6, 1.6],
        style=TableStyle(alternate_row_background_color="#EEF4FA"),
    )
    merged_header_table = Table(
        headers=[
            [TableCell("Metrics", colspan=2), TableCell("Summary", rowspan=2)],
            ["Latency", "Quality"],
        ],
        rows=[
            [
                TableCell("14 ms"),
                TableCell("stable", background_color="#EEF6FF"),
                TableCell("ready"),
            ]
        ],
        caption="Merged header table.",
        column_widths=[1.6, 1.6, 1.6],
        style=TableStyle(header_background_color="#D9E6F2"),
    )
    preview_figure = Figure(
        image_path,
        caption=Paragraph("Tiny sample image."),
        width_inches=1.0,
    )
    figure_object = FakeFigure(_build_sample_png(width=320, height=180))
    preview_figure_second = Figure(
        figure_object,
        caption=Paragraph("Second tiny sample image."),
        width_inches=1.2,
    )
    boxed_detail = Box(
        Paragraph("A boxed paragraph can live alongside nested objects."),
        Table(
            headers=["Scope", "State"],
            rows=[["Box", "stable"]],
            column_widths=[1.4, 1.4],
        ),
        Figure(
            image_path,
            width_inches=0.7,
        ),
        title="Review Box",
        style=BoxStyle(
            border_color="#7A8CA5",
            background_color="#F4F8FC",
            title_background_color="#DDE8F4",
        ),
    )

    document = Document(
        "Pipeline Report",
        TableOfContents(),
        Chapter(
            "Summary",
            Section(
                "Highlights",
                Paragraph(
                    "The review ",
                    comment("note", "Check the generated outputs before release.", author="pytest", initials="PT"),
                    " appears inline and is also exported to the comments page.",
                ),
                HighlightedParagraph(
                    "The ",
                    bold("docscriptor"),
                    " pipeline supports ",
                    italic("styled"),
                    " text, ",
                    code("code"),
                    ", and ",
                        color("custom color", "#0066AA", style=TextStyle(bold=True)),
                    ".",
                    style=ParagraphStyle(space_after=14),
                ),
                Paragraph(
                    markup("Inline helpers also support **bold** and *italic* markup."),
                    " Inline math such as ",
                    math(r"\alpha^2 + \beta^2 = \gamma^2"),
                    " is supported as well.",
                ),
                Paragraph(
                    "See ",
                    artifacts_table,
                    " and ",
                    preview_figure,
                    " for the generated outputs.",
                ),
                Paragraph(
                    "Repository status is tracked in ",
                    cite(repository_source),
                    ".",
                ),
                Paragraph(
                    "Registered bibliography entries can still be cited as ",
                    cite("release-notes"),
                    ".",
                ),
                Paragraph(
                    "Portable footnotes such as ",
                    footnote("term", "Paragraph footnote note."),
                    " are collected automatically on the footnotes page.",
                ),
                boxed_detail,
                Subsection(
                    "Artifacts",
                    BulletList(
                        "Lists render into both DOCX and PDF.",
                        Paragraph("Paragraph instances can also be list items."),
                    ),
                    Subsubsection(
                        "Export Steps",
                        CodeBlock(
                            "from docscriptor import Document\n\ndocument.save_docx('report.docx')\ndocument.save_pdf('report.pdf')",
                            language="python",
                        ),
                    ),
                    Equation(r"\int_0^1 \alpha x^2 \, dx = \frac{\alpha}{3}"),
                    NumberedList("Create the model", "Render the files"),
                    artifacts_table,
                    workflow_table,
                    merged_header_table,
                    preview_figure,
                    preview_figure_second,
                ),
            ),
        ),
        TableList(),
        FigureList(),
        CommentsPage(),
        ReferencesPage(),
        settings=DocumentSettings(
            author="pytest",
            summary="Renderer integration test",
            theme=Theme(
                show_page_numbers=True,
                page_number_format="Page {page}",
                page_number_alignment="center",
                footnote_placement="document",
                heading_numbering=HeadingNumbering(),
                bullet_list_style=ListStyle(marker_format="bullet", bullet="\u2022", suffix=""),
                numbered_list_style=ListStyle(marker_format="decimal", suffix="."),
            ),
        ),
        citations=[registered_source, unused_source],
    )

    docx_path = tmp_path / "report.docx"
    pdf_path = tmp_path / "report.pdf"
    html_path = tmp_path / "report.html"

    document.save_docx(docx_path)
    document.save_pdf(pdf_path)
    document.save_html(html_path)

    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert docx_path.stat().st_size > 0
    assert pdf_path.stat().st_size > 0
    assert html_path.stat().st_size > 0

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    assert "Pipeline Report" in paragraph_texts
    assert "1 Summary" in paragraph_texts
    assert "1.1 Highlights" in paragraph_texts
    assert "1.1.1 Artifacts" in paragraph_texts
    assert "1.1.1.1 Export Steps" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "Comments" in paragraph_texts
    assert "Footnotes" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "List of Figures" in paragraph_texts
    assert "References" in paragraph_texts
    assert any("The review note[1] appears inline" in text for text in paragraph_texts)
    assert any("docscriptor" in text for text in paragraph_texts)
    assert any(text == "1 Summary" for text in paragraph_texts)
    assert any(text == "1.1 Highlights" for text in paragraph_texts)
    assert any("See Table 1 and Figure 1 for the generated outputs." in text for text in paragraph_texts)
    assert any("Repository status is tracked in [1]." in text for text in paragraph_texts)
    assert any("Registered bibliography entries can still be cited as [2]." in text for text in paragraph_texts)
    assert any("Portable footnotes such as term" in text and "collected automatically on the footnotes page." in text for text in paragraph_texts)
    assert any("Inline math such as" in text and "2 + " in text and " = " in text for text in paragraph_texts)
    assert any("dx = (" in text and ")/(3)" in text for text in paragraph_texts)
    assert any("Table cell footnote note." in text for text in paragraph_texts)
    assert any("Paragraph footnote note." in text for text in paragraph_texts)
    assert any("[1] Check the generated outputs before release." in text for text in paragraph_texts)
    assert any(text == "• Lists render into both DOCX and PDF." for text in paragraph_texts)
    assert any(text == "1. Create the model" for text in paragraph_texts)
    assert paragraph_texts.count("Table 1. Generated artifacts.") >= 2
    assert paragraph_texts.count("Table 2. Output workflow.") >= 2
    assert paragraph_texts.count("Table 3. Merged header table.") >= 2
    assert paragraph_texts.count("Figure 1. Tiny sample image.") >= 2
    assert paragraph_texts.count("Figure 2. Second tiny sample image.") >= 2
    assert any("https://github.com/Gonie-Gonie/pydocs" in text for text in paragraph_texts)
    assert any("https://github.com/Gonie-Gonie/pydocs/releases" in text for text in paragraph_texts)
    assert all("internal-draft" not in text.lower() for text in paragraph_texts)
    assert any("from docscriptor import Document" in text for text in paragraph_texts)
    assert len(word_document.inline_shapes) == 3
    summary_paragraph = next(
        paragraph
        for paragraph in word_document.paragraphs
        if "The review note[1] appears inline" in paragraph.text
    )
    assert summary_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    assert len(word_document.tables) == 4
    assert "Review Box" in word_document.tables[0].cell(0, 0).text
    assert word_document.tables[1].cell(1, 0).text == "DOCX"
    assert word_document.tables[1].cell(1, 1).text.startswith("generated")
    assert word_document.tables[1].cell(2, 1).text == "generated"
    assert word_document.tables[2].cell(2, 0).text == "Draft review"
    assert word_document.tables[2].cell(3, 1).text == "PDF"
    assert word_document.tables[3].cell(2, 0).text == "14 ms"
    assert word_document.tables[3].cell(2, 2).text == "ready"
    assert word_document.styles["Normal"].font.name == "Times New Roman"
    assert word_document.styles["Title"].font.name == "Times New Roman"
    assert word_document.styles["Heading 1"].font.name == "Times New Roman"
    assert word_document.styles["Heading 2"].font.name == "Times New Roman"
    assert word_document.styles["Heading 3"].font.name == "Times New Roman"
    assert word_document.styles["Heading 4"].font.name == "Times New Roman"
    assert word_document.styles["Title"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 1"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 2"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 3"].font.color.rgb == RGBColor(0, 0, 0)
    assert word_document.styles["Heading 4"].font.color.rgb == RGBColor(0, 0, 0)
    heading_styles = {paragraph.text: paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text in {"Comments", "List of Tables", "List of Figures", "References"}}
    assert heading_styles["Comments"] == "Heading 2"
    assert heading_styles["List of Tables"] == "Heading 2"
    assert heading_styles["List of Figures"] == "Heading 2"
    assert heading_styles["References"] == "Heading 2"
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "Footnotes") == "Heading 2"
    assert len(word_document.comments) == 1
    assert "Check the generated outputs before release." in "\n".join(
        paragraph.text
        for comment_item in word_document.comments
        for paragraph in comment_item.paragraphs
    )
    footer_xml = word_document.sections[0].footer.paragraphs[0]._p.xml
    assert 'w:instr="PAGE"' in footer_xml
    assert word_document.sections[0].footer.paragraphs[0].text.startswith("Page ")
    inline_math_paragraph = next(paragraph for paragraph in word_document.paragraphs if "Inline math such as" in paragraph.text)
    assert any(run.text == "2" and run.font.superscript for run in inline_math_paragraph.runs)
    equation_paragraph = next(paragraph for paragraph in word_document.paragraphs if "dx = (" in paragraph.text and ")/(3)" in paragraph.text)
    assert any(run.text == "2" and run.font.superscript for run in equation_paragraph.runs)
    assert any(run.text == "0" and run.font.subscript for run in equation_paragraph.runs)
    assert any(run.text == "1" and run.font.superscript for run in equation_paragraph.runs)
    docx_xml = _docx_document_xml(docx_path)
    assert "Review Box" in docx_xml
    assert "A boxed paragraph can live alongside nested objects." in docx_xml
    assert "stable" in docx_xml
    assert "D9E6F2" in docx_xml
    assert "DCE8F4" in docx_xml
    assert len(figure_object.calls) >= 2
    assert all(call.get("format") == "png" for call in figure_object.calls)

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages)
    assert "Pipeline Report" in pdf_text
    assert "1 Summary" in pdf_text
    assert "1.1 Highlights" in pdf_text
    assert "1.1.1 Artifacts" in pdf_text
    assert "1.1.1.1 Export Steps" in pdf_text
    assert "Contents" in pdf_text
    assert "Comments" in pdf_text
    assert "Footnotes" in pdf_text
    assert "The review note[1] appears inline and is also exported to the comments page." in pdf_text
    assert "See Table 1 and Figure 1 for the generated outputs." in pdf_text
    assert "Repository status is tracked in [1]." in pdf_text
    assert "Registered bibliography entries can still be cited as [2]." in pdf_text
    assert "Portable footnotes such as term" in pdf_text and "collected automatically on the footnotes page." in pdf_text
    assert "Inline math such as" in pdf_text
    assert "dx = (" in pdf_text
    assert "Table cell footnote note." in pdf_text
    assert "Paragraph footnote note." in pdf_text
    assert "[1] Check the generated outputs before release." in pdf_text
    assert "Review Box" in pdf_text
    assert "A boxed paragraph can live alongside nested objects." in pdf_text
    assert "stable" in pdf_text
    assert pdf_text.count("Table 1. Generated artifacts.") >= 2
    assert pdf_text.count("Table 2. Output workflow.") >= 2
    assert pdf_text.count("Table 3. Merged header table.") >= 2
    assert pdf_text.count("Figure 1. Tiny sample image.") >= 2
    assert pdf_text.count("Figure 2. Second tiny sample image.") >= 2
    assert "List of Tables" in pdf_text
    assert "List of Figures" in pdf_text
    assert "References" in pdf_text
    assert "https://github.com/Gonie-Gonie/pydocs" in pdf_text
    assert "https://github.com/Gonie-Gonie/pydocs/releases" in pdf_text
    assert "Internal Draft" not in pdf_text
    assert "Lists render into both DOCX and PDF." in pdf_text
    assert "from docscriptor import Document" in pdf_text
    assert "1\nLists render into both DOCX and PDF." not in pdf_text
    assert "1.\nCreate the model" in pdf_text
    assert _pdf_image_draw_count(pdf_path) == 3
    pdf_fonts = _pdf_font_names(pdf_path)
    assert any(font == "/Times-Roman" or "TimesNewRomanPSMT" in font for font in pdf_fonts)
    assert any(font == "/Times-Bold" or "TimesNewRomanPS-Bold" in font for font in pdf_fonts)
    assert any(
        font in {"/Times-BoldItalic", "/Times-Italic"}
        or "TimesNewRomanPS-BoldItalic" in font
        or "TimesNewRomanPS-Italic" in font
        for font in pdf_fonts
    )
    assert any(font.startswith("/Courier") or "CourierNewPS" in font for font in pdf_fonts)
    pdf_content = _pdf_content_bytes(pdf_path)
    assert b"Page 1" in pdf_content
    assert b"18 Tf" in pdf_content
    assert b"15 Tf" in pdf_content
    assert b"13 Tf" in pdf_content
    assert b"11.5 Tf" in pdf_content
    assert b"Comments" in pdf_content
    assert b"15 Tf" in _pdf_text_context(pdf_path, "List of Tables")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "List of Figures")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "Comments")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "References")
    assert b"15 Tf" in _pdf_text_context(pdf_path, "Contents")

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "Pipeline Report" in normalized_html_text
    assert "1 Summary" in normalized_html_text
    assert "1.1 Highlights" in normalized_html_text
    assert "1.1.1 Artifacts" in normalized_html_text
    assert "1.1.1.1 Export Steps" in normalized_html_text
    assert "Contents" in normalized_html_text
    assert "Comments" in normalized_html_text
    assert "Footnotes" in normalized_html_text
    assert "List of Tables" in normalized_html_text
    assert "List of Figures" in normalized_html_text
    assert "References" in normalized_html_text
    assert "The review note [1] appears inline and is also exported to the comments page." in normalized_html_text
    assert "See Table 1 and Figure 1 for the generated outputs." in normalized_html_text
    assert "Repository status is tracked in [1]" in normalized_html_text
    assert "Registered bibliography entries can still be cited as [2]" in normalized_html_text
    assert "Portable footnotes such as term" in normalized_html_text
    assert "Table cell footnote note." in normalized_html_text
    assert "Paragraph footnote note." in normalized_html_text
    assert "[1] Check the generated outputs before release." in normalized_html_text
    assert "Review Box" in normalized_html_text
    assert "A boxed paragraph can live alongside nested objects." in normalized_html_text
    assert "stable" in normalized_html_text
    assert normalized_html_text.count("Table 1. Generated artifacts.") >= 2
    assert normalized_html_text.count("Table 2. Output workflow.") >= 2
    assert normalized_html_text.count("Table 3. Merged header table.") >= 2
    assert normalized_html_text.count("Figure 1. Tiny sample image.") >= 2
    assert normalized_html_text.count("Figure 2. Second tiny sample image.") >= 2
    assert "https://github.com/Gonie-Gonie/pydocs" in normalized_html_text
    assert "https://github.com/Gonie-Gonie/pydocs/releases" in normalized_html_text
    assert "Internal Draft" not in normalized_html_text
    assert "Lists render into both DOCX and PDF." in normalized_html_text
    assert "from docscriptor import Document" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 3
    assert 'href="#table_1"' in html_text
    assert 'href="#figure_1"' in html_text
    assert 'id="table_1"' in html_text
    assert 'id="figure_1"' in html_text
    assert 'id="citation_1"' in html_text
    assert 'id="comment_1"' in html_text
    assert 'id="footnote_1"' in html_text
