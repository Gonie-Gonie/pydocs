from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re
import zipfile

from docx import Document as WordDocument
from pypdf import PdfReader


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_usage_guide_example_builds_outputs(tmp_path: Path) -> None:
    usage_guide = _load_example_module("usage_guide_example")
    docx_path, pdf_path = usage_guide.build_usage_guide(tmp_path)
    html_path = tmp_path / "docscriptor-user-guide.html"

    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert (Path(usage_guide.__file__).resolve().parent / "assets" / "docscriptor-logo.png").exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)

    assert "Docscriptor User Guide" in paragraph_texts
    assert "Reference-style guide for structured Python document authoring" in paragraph_texts
    assert "Docscriptor Contributors" in paragraph_texts
    assert "Open-source documentation workflow" in paragraph_texts
    assert "Hyeong-Gon Jo" in paragraph_texts
    assert "Repository steward" in paragraph_texts
    assert "Guide Cover" in paragraph_texts
    assert "Contents" in paragraph_texts
    assert "List of Tables" in paragraph_texts
    assert "List of Figures" in paragraph_texts
    assert "Comments" in paragraph_texts
    assert "References" in paragraph_texts
    assert "1 Overview" in paragraph_texts
    assert "2 Metadata and Title Matter" in paragraph_texts
    assert "3 Document Model" in paragraph_texts
    assert "4 Tables, Figures, and Cross-References" in paragraph_texts
    assert "5 Notes, Comments, and References" in paragraph_texts
    assert "6 Layout and Pagination" in paragraph_texts
    assert "7 Project Structure and Scaling Up" in paragraph_texts
    assert any("AuthorLayout(mode='stacked')" in text for text in paragraph_texts)
    assert any("Theme(footnote_placement='document')" in text for text in paragraph_texts)
    assert any("PageMargins" in text for text in paragraph_texts)
    assert any("PageBreak()" in text for text in paragraph_texts)
    assert any("settings.get_text_width(0.75)" in text for text in paragraph_texts)
    assert any("A reading map for the guide." in text for text in paragraph_texts)
    assert any("Page layout controls shared across renderers." in text for text in paragraph_texts)
    assert any("Figure sizing patterns for width, height, and document-relative sizing." in text for text in paragraph_texts)
    assert any("Renderer-specific behavior for notes, review workflows, and cross-reference stability." in text for text in paragraph_texts)
    assert any("portable footnotes exactly where the text appears." in text for text in paragraph_texts)
    assert any("github.com/Gonie-Gonie/docscriptor" in text for text in paragraph_texts)
    assert any("The journal example at examples/journal_paper_example/main.py" in text for text in paragraph_texts)
    assert "Footnotes" not in [text for text in paragraph_texts if text == "Footnotes"]
    assert len(word_document.tables) == 9
    assert len(word_document.inline_shapes) == 4
    assert len(word_document.comments) == 2
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "Comments") == "Heading 2"
    assert next(paragraph.style.name for paragraph in word_document.paragraphs if paragraph.text == "References") == "Heading 2"

    with zipfile.ZipFile(docx_path) as archive:
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
    assert "DOCX uses page footnotes by default." in footnotes_xml
    assert "Portable footnotes are authored inline" in footnotes_xml
    assert all("CommentsPage() collects these review notes onto a dedicated generated page." in "\n".join(p.text for p in comment.paragraphs) or "This note will show up again on the generated comments page." in "\n".join(p.text for p in comment.paragraphs) for comment in word_document.comments)

    assert "Docscriptor User Guide" in pdf_text
    assert "Contents" in pdf_text
    assert "List of Tables" in pdf_text
    assert "List of Figures" in pdf_text
    assert "Comments" in pdf_text
    assert "References" in pdf_text
    assert "Guide Cover" in pdf_text
    assert "AuthorLayout(mode='stacked')" in pdf_text
    assert "Theme(footnote_placement='document')" in pdf_text
    assert "PageMargins" in pdf_text
    assert "PageBreak()" in pdf_text
    assert "settings.get_text_width(0.75)" in pdf_text
    assert "A reading map for the guide." in pdf_text
    assert "Page layout controls shared across renderers." in pdf_text
    assert "Figure sizing patterns for width, height, and document-relative sizing." in pdf_text
    assert "Renderer-specific behavior for notes, review workflows, and cross-reference stability." in pdf_text
    assert "Portable footnotes are authored inline" in pdf_text
    assert "github.com/Gonie-Gonie/docscriptor" in pdf_text
    assert "Footnotes" in pdf_text
    assert len(pdf_reader.pages) >= 14
    assert _pdf_image_draw_count(pdf_path) == 4

    assert "Docscriptor User Guide" in normalized_html_text
    assert "Guide Cover" in normalized_html_text
    assert "List of Tables" in normalized_html_text
    assert "List of Figures" in normalized_html_text
    assert "Comments" in normalized_html_text
    assert "References" in normalized_html_text
    assert "AuthorLayout(mode='stacked')" in normalized_html_text
    assert "Theme(footnote_placement='document')" in normalized_html_text
    assert "PageMargins" in normalized_html_text
    assert "PageBreak()" in normalized_html_text
    assert "settings.get_text_width(0.75)" in normalized_html_text
    assert "CommentsPage() collects these review notes onto a dedicated generated page." in normalized_html_text
    assert "Page layout controls shared across renderers." in normalized_html_text
    assert "Figure sizing patterns for width, height, and document-relative sizing." in normalized_html_text
    assert "Portable footnotes are authored inline" in normalized_html_text
    assert "github.com/Gonie-Gonie/docscriptor" in normalized_html_text
    assert "Footnotes" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 4
    assert 'href="#table_1"' in html_text
    assert 'href="#figure_1"' in html_text
    assert 'class="docscriptor-toc-entry docscriptor-toc-entry-level-1"' in html_text
    assert 'class="docscriptor-toc-entry docscriptor-toc-entry-level-2"' in html_text
