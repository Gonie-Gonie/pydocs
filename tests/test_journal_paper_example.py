from __future__ import annotations

import importlib.util
from html import unescape
from io import BytesIO
from pathlib import Path
import re

from docx import Document as WordDocument
from pypdf import PdfReader


def _load_example_module(example_dir: str):
    module_path = Path(__file__).resolve().parents[1] / "examples" / example_dir / "main.py"
    spec = importlib.util.spec_from_file_location(f"examples.{example_dir}.main", module_path)
    assert spec is not None
    assert spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _pdf_image_draw_count(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(BytesIO(pdf_path.read_bytes())).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        image_names = {
            name
            for name, xobject in xobjects.items()
            if xobject.get_object().get("/Subtype") == "/Image"
        }
        if not image_names:
            continue
        content = page.get_contents()
        if content is None:
            continue
        content_bytes = content.get_data()
        for name in image_names:
            token = f"{name} Do".encode()
            count += content_bytes.count(token)
    return count


def _normalized_html_text(html_path: Path) -> str:
    html_text = html_path.read_text(encoding="utf-8")
    html_text = re.sub(r"<style.*?>.*?</style>", " ", html_text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", " ", html_text)
    return " ".join(unescape(text).split())


def test_journal_paper_example_builds_outputs(tmp_path: Path) -> None:
    paper_example = _load_example_module("journal_paper_example")
    docx_path, pdf_path = paper_example.build_journal_paper(tmp_path)
    html_path = tmp_path / "journal-paper.html"

    assert docx_path.exists()
    assert pdf_path.exists()
    assert html_path.exists()
    assert (Path(paper_example.__file__).resolve().parent / "assets" / "benchmark_results.csv").exists()
    assert (Path(paper_example.__file__).resolve().parent / "assets" / "ablation_results.csv").exists()
    assert (Path(paper_example.__file__).resolve().parent / "assets" / "system-diagram.png").exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    pdf_reader = PdfReader(BytesIO(pdf_path.read_bytes()))
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    normalized_pdf_text = " ".join(pdf_text.split())

    assert "A Python-Native Workflow for Reproducible Journal Manuscripts" in paragraph_texts
    assert "Jiyoon Kim, Minho Lee, and Sujin Park" in paragraph_texts
    assert "Department of Computational Publishing, Seoul" in paragraph_texts
    assert "Abstract" in paragraph_texts
    assert "Highlights" in paragraph_texts
    assert "Acknowledgements" in paragraph_texts
    assert "1 Introduction" in paragraph_texts
    assert "2 Related Work" in paragraph_texts
    assert "3 Methods" in paragraph_texts
    assert "4 Experimental Setup" in paragraph_texts
    assert "5 Results" in paragraph_texts
    assert "7 Discussion" in paragraph_texts
    assert "8 Conclusion" in paragraph_texts
    assert "References" in paragraph_texts
    assert "3.1 Asset Layout" in paragraph_texts
    assert "3.2 Data Integration" in paragraph_texts
    assert "5.1 Benchmark Performance" in paragraph_texts
    assert "5.2 Ablation Study" in paragraph_texts
    assert all(text not in paragraph_texts for text in ("Contents", "List of Tables", "List of Figures"))
    assert any("pandas.read_csv" in text for text in paragraph_texts)
    assert any("matplotlib" in text for text in paragraph_texts)
    assert any("benchmark CSV" in text for text in paragraph_texts)
    assert any("As summarized in Table 2" in text and "Figure 2" in text for text in paragraph_texts)
    assert paragraph_texts.count("Table 1. Corpus summary used for the manuscript automation study.") >= 1
    assert paragraph_texts.count("Table 2. Benchmark results loaded directly from the experiment CSV file.") >= 1
    assert paragraph_texts.count("Table 3. Ablation results for major document-pipeline design decisions.") >= 1
    assert paragraph_texts.count("Figure 1. System overview diagram stored under the example asset directory.") >= 1
    assert paragraph_texts.count("Figure 2. Accuracy and latency curves generated directly from the benchmark CSV with matplotlib.") >= 1
    assert paragraph_texts.count("Figure 3. Ablation accuracy chart generated from the ablation CSV with matplotlib.") >= 1
    assert len(word_document.tables) == 3
    assert len(word_document.inline_shapes) == 3

    assert "Journal Manuscripts" in pdf_text
    assert "Abstract" in pdf_text
    assert "Highlights" in pdf_text
    assert "Introduction" in pdf_text
    assert "Methods" in pdf_text
    assert "Results" in pdf_text
    assert "Discussion" in pdf_text
    assert "Acknowledgements" in pdf_text
    assert "References" in pdf_text
    assert "Asset Layout" in pdf_text
    assert "Data Integration" in pdf_text
    assert "Benchmark Performance" in pdf_text
    assert "Ablation Study" in pdf_text
    assert "Contents" not in pdf_text
    assert "List of Tables" not in pdf_text
    assert "List of Figures" not in pdf_text
    assert "pandas.read_csv" in pdf_text
    assert "matplotlib" in pdf_text
    assert "benchmark CSV" in pdf_text
    assert "Literate Programming" in normalized_pdf_text
    assert "Statistical Analyses and Reproducible Research" in normalized_pdf_text
    assert "https://doi.org/10.1093/comjnl/27.2.97" in normalized_pdf_text
    assert "https://doi.org/10.1198/106186007X178663" in normalized_pdf_text
    assert "https://yihui.org/knitr/" in normalized_pdf_text
    assert 6 <= len(pdf_reader.pages) <= 8
    assert _pdf_image_draw_count(pdf_path) == 3

    html_text = html_path.read_text(encoding="utf-8")
    normalized_html_text = _normalized_html_text(html_path)
    assert "A Python-Native Workflow for Reproducible Journal Manuscripts" in normalized_html_text
    assert "Jiyoon Kim, Minho Lee, and Sujin Park" in normalized_html_text
    assert "Department of Computational Publishing, Seoul" in normalized_html_text
    assert "Abstract" in normalized_html_text
    assert "Highlights" in normalized_html_text
    assert "Introduction" in normalized_html_text
    assert "Methods" in normalized_html_text
    assert "Experimental Setup" in normalized_html_text
    assert "Results" in normalized_html_text
    assert "Discussion" in normalized_html_text
    assert "Acknowledgements" in normalized_html_text
    assert "References" in normalized_html_text
    assert "DOCX, PDF, and HTML outputs are rendered from the same source document." in normalized_html_text
    assert "The same manuscript source renders to DOCX, PDF, and HTML for review and submission." in normalized_html_text
    assert "pandas.read_csv" in normalized_html_text
    assert "matplotlib" in normalized_html_text
    assert "https://doi.org/10.1093/comjnl/27.2.97" in normalized_html_text
    assert "https://doi.org/10.1198/106186007X178663" in normalized_html_text
    assert "https://yihui.org/knitr/" in normalized_html_text
    assert html_text.count("data:image/png;base64,") == 3
    assert 'href="#table_2"' in html_text
    assert 'href="#figure_2"' in html_text
