from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader

from examples.showcase import SHOWCASE_IMAGE, build_showcase


def _count_pdf_images(pdf_path: Path) -> int:
    count = 0
    for page in PdfReader(str(pdf_path)).pages:
        resources = page.get("/Resources")
        if resources is None or "/XObject" not in resources:
            continue
        xobjects = resources["/XObject"].get_object()
        for xobject in xobjects.values():
            obj = xobject.get_object()
            if obj.get("/Subtype") == "/Image":
                count += 1
    return count


def test_showcase_script_builds_outputs(tmp_path: Path) -> None:
    docx_path, pdf_path = build_showcase(tmp_path)

    assert docx_path.exists()
    assert pdf_path.exists()
    assert SHOWCASE_IMAGE.exists()

    word_document = WordDocument(docx_path)
    paragraph_texts = [paragraph.text for paragraph in word_document.paragraphs]
    assert "Docscriptor Showcase" in paragraph_texts
    assert "Feature Checklist" in paragraph_texts
    assert any(paragraph.style.name == "List Bullet" for paragraph in word_document.paragraphs)
    assert any(paragraph.style.name == "List Number" for paragraph in word_document.paragraphs)
    assert any("Bundled showcase image." in text for text in paragraph_texts)
    assert len(word_document.inline_shapes) >= 1

    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)
    assert "Docscriptor Showcase" in pdf_text
    assert "Authoring Model" in pdf_text
    assert "Current showcase capabilities." in pdf_text
    assert "Sections and subsections are regular Python instances." in pdf_text
    assert "Bundled showcase image." in pdf_text
    assert _count_pdf_images(pdf_path) >= 1
