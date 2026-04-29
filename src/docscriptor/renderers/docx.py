"""DOCX renderer."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_BREAK
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.parts.story import StoryPart
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph

from docscriptor.components.blocks import (
    Box,
    BulletList,
    CodeBlock,
    Equation,
    NumberedList,
    PageBreak,
    Paragraph,
    Section,
)
from docscriptor.components.generated import (
    CommentsPage,
    FigureList,
    FootnotesPage,
    ReferencesPage,
    TableList,
    TableOfContents,
    TocLevelStyle,
)
from docscriptor.components.inline import (
    _BlockReference,
    Citation,
    Comment,
    Footnote,
    Hyperlink,
    Math,
    Text,
)
from docscriptor.components.media import Figure, Table, build_table_layout
from docscriptor.components.people import AuthorTitleLine
from docscriptor.components.sheets import ImageBox, Sheet, TextBox
from docscriptor.document import Document
from docscriptor.components.equations import SUBSCRIPT, SUPERSCRIPT, parse_latex_segments
from docscriptor.core import DocscriptorError, PathLike, length_to_inches
from docscriptor.layout.indexing import RenderIndex, build_render_index
from docscriptor.layout.theme import ParagraphStyle, TextStyle, Theme
from docscriptor.renderers.context import DocxRenderContext


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

TABLE_ALIGNMENTS = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}

DEFAULT_FOOTNOTES_XML = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator />
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator />
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>
"""


class FootnotesPart(StoryPart):
    """Container part for native DOCX footnotes."""

    @classmethod
    def default(cls, package: object) -> "FootnotesPart":
        return cls(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            parse_xml(DEFAULT_FOOTNOTES_XML),
            package,
        )

    def add_footnote_paragraph(self, footnote_id: int) -> DocxParagraph:
        footnote = OxmlElement("w:footnote")
        footnote.set(qn("w:id"), str(footnote_id))
        paragraph = OxmlElement("w:p")
        footnote.append(paragraph)
        self._element.append(footnote)
        return DocxParagraph(paragraph, self)


class DocxRenderer:
    """Render docscriptor documents into DOCX files."""

    def render(self, document: Document, output_path: PathLike) -> Path:
        """Render a docscriptor document to a DOCX file."""

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        word_document = WordDocument()
        self._initialized_cells: set[int] = set()
        self._bookmark_id = 1
        self._native_footnotes_part: FootnotesPart | None = None
        self._rendered_native_footnotes: set[int] = set()
        render_index = build_render_index(document)
        self._configure_document(word_document, document)
        context = DocxRenderContext(
            theme=document.theme,
            render_index=render_index,
            settings=document.settings,
            unit=document.settings.unit,
            word_document=word_document,
        )
        front_children, main_children = document.split_top_level_children()
        has_front_matter = document.cover_page or bool(front_children)

        self._render_title_matter(
            word_document,
            document,
            context,
        )

        if document.cover_page and front_children:
            self._ensure_page_break(word_document)

        if has_front_matter:
            self._render_top_level_children(word_document, front_children, context)
            if main_children:
                section = word_document.add_section(WD_SECTION.NEW_PAGE)
                self._configure_section_page_box(section, document.settings)
                self._render_top_level_children(word_document, main_children, context)
        else:
            self._render_top_level_children(word_document, main_children, context)

        if self._should_auto_render_footnotes_page(document, render_index):
            self.render_footnotes_page(FootnotesPage(), context)

        if document.theme.show_page_numbers:
            self._configure_page_number_sections(
                word_document,
                document.theme,
                has_front_matter=has_front_matter,
                has_main_matter=bool(main_children) or not has_front_matter,
            )

        word_document.save(path)
        return path

    def add_heading(
        self,
        container: object,
        title: list[Text],
        level: int,
        context: DocxRenderContext,
        *,
        number_label: str | None = None,
        anchor: str | None = None,
    ) -> None:
        """Render a heading into the current DOCX container."""

        self._add_heading(
            container,
            title,
            level,
            context.theme,
            number_label=number_label,
            anchor=anchor,
        )

    def render_paragraph(
        self,
        container: object,
        paragraph_block: Paragraph,
        context: DocxRenderContext,
    ) -> None:
        """Render a paragraph block into the current DOCX container."""

        paragraph = self._add_paragraph(container)
        self._apply_paragraph_style(paragraph, paragraph_block.style)
        self._append_runs(
            paragraph,
            paragraph_block.content,
            theme=context.theme,
            render_index=context.render_index,
            word_document=context.word_document,
        )

    def render_list(
        self,
        container: object,
        list_block: BulletList | NumberedList,
        context: DocxRenderContext,
    ) -> None:
        """Render a list block into the current DOCX container."""

        self._render_list(
            container,
            list_block,
            context.theme,
            context.render_index,
            word_document=context.word_document,
        )

    def render_code_block(
        self,
        container: object,
        code_block: CodeBlock,
        context: DocxRenderContext,
    ) -> None:
        """Render a code block into the current DOCX container."""

        self._render_code_block(container, code_block, context.theme)

    def render_equation(
        self,
        container: object,
        equation: Equation,
        context: DocxRenderContext,
    ) -> None:
        """Render a block equation into the current DOCX container."""

        self._render_equation(container, equation, context.theme)

    def render_page_break(
        self,
        container: object,
        block: PageBreak,
        context: DocxRenderContext,
    ) -> None:
        """Render an explicit page break into the current DOCX container."""

        paragraph = self._add_paragraph(container)
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def render_box(
        self,
        container: object,
        box: Box,
        context: DocxRenderContext,
    ) -> None:
        """Render a box and its child blocks into the current DOCX container."""

        self._render_box(
            container,
            box,
            context.theme,
            context.render_index,
            context.settings,
            context.unit,
            word_document=context.word_document,
        )

    def render_sheet(
        self,
        container: object,
        sheet: Sheet,
        context: DocxRenderContext,
    ) -> None:
        """Render a fixed-layout sheet into the current DOCX container."""

        self._render_sheet(
            container,
            sheet,
            context.theme,
            context.render_index,
            context.settings,
            context.unit,
            word_document=context.word_document,
        )
        if sheet.page_break_after:
            self._ensure_page_break(context.word_document)

    def render_comments_page(
        self,
        block: CommentsPage,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated comments page into the DOCX document."""

        self._render_comments_page(
            context.word_document,
            block.title,
            context.theme,
            context.render_index,
        )

    def render_footnotes_page(
        self,
        block: FootnotesPage,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated footnotes page into the DOCX document."""

        self._render_footnotes_page(
            context.word_document,
            block.title,
            context.theme,
            context.render_index,
        )

    def render_references_page(
        self,
        block: ReferencesPage,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated references page into the DOCX document."""

        self._render_references_page(
            context.word_document,
            block.title,
            context.theme,
            context.render_index,
        )

    def render_table_of_contents(
        self,
        block: TableOfContents,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated table of contents into the DOCX document."""

        self._render_table_of_contents(
            context.word_document,
            block,
            context,
        )

    def render_table_list(
        self,
        block: TableList,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of tables into the DOCX document."""

        self._render_caption_list(
            context.word_document,
            block.title,
            context.render_index.tables,
            context.theme,
            context.render_index,
            context.theme.list_of_tables_title,
            context.theme.table_label,
        )

    def render_figure_list(
        self,
        block: FigureList,
        context: DocxRenderContext,
    ) -> None:
        """Render the generated list of figures into the DOCX document."""

        self._render_caption_list(
            context.word_document,
            block.title,
            context.render_index.figures,
            context.theme,
            context.render_index,
            context.theme.list_of_figures_title,
            context.theme.figure_label,
        )

    def render_table(
        self,
        container: object,
        table_block: Table,
        context: DocxRenderContext,
    ) -> None:
        """Render a table block into the current DOCX container."""

        self._render_table(
            container,
            table_block,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
        )

    def render_figure(
        self,
        container: object,
        figure: Figure,
        context: DocxRenderContext,
    ) -> None:
        """Render a figure block into the current DOCX container."""

        self._render_figure(
            container,
            figure,
            context.theme,
            context.render_index,
            context.unit,
            word_document=context.word_document,
        )

    def _configure_document(self, word_document: WordDocument, document: Document) -> None:
        properties = word_document.core_properties
        properties.title = document.title
        if document.author:
            properties.author = document.author
        if document.summary:
            properties.subject = document.summary

        normal_style = word_document.styles["Normal"]
        normal_style.font.name = document.theme.body_font_name
        normal_style.font.size = Pt(document.theme.body_font_size)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        footer_style = word_document.styles["Footer"]
        footer_style.font.name = document.theme.body_font_name
        footer_style.font.size = Pt(document.theme.page_number_font_size)
        footer_style.font.color.rgb = RGBColor(0, 0, 0)
        for section in word_document.sections:
            self._configure_section_page_box(section, document.settings)

        self._configure_named_style(
            word_document,
            "Title",
            font_name=document.theme.body_font_name,
            font_size=document.theme.title_font_size,
            bold=True,
            italic=False,
        )
        for level in range(1, 5):
            bold, italic = document.theme.heading_emphasis(level)
            self._configure_named_style(
                word_document,
                f"Heading {level}",
                font_name=document.theme.body_font_name,
                font_size=document.theme.heading_size(level),
                bold=bold,
                italic=italic,
            )
        self._set_page_background(word_document, document.theme.page_background_color)

    def _render_block(
        self,
        container: object,
        block: object,
        context: DocxRenderContext,
    ) -> None:
        """Delegate block rendering back to the block instance itself."""

        block.render_to_docx(self, container, context)

    def _configure_section_page_box(self, section: object, settings: object) -> None:
        top, right, bottom, left = settings.page_margin_inches()
        section.page_width = Inches(settings.page_width_in_inches())
        section.page_height = Inches(settings.page_height_in_inches())
        section.top_margin = Inches(top)
        section.right_margin = Inches(right)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)

    def _render_top_level_children(
        self,
        word_document: WordDocument,
        children: list[object],
        context: DocxRenderContext,
    ) -> None:
        for index, child in enumerate(children):
            if self._is_paginated_generated_page(child) and context.theme.generated_page_breaks:
                if word_document.paragraphs and not self._ends_with_page_break(word_document):
                    self._ensure_page_break(word_document)
                child.render_to_docx(self, word_document, context)
                if index < len(children) - 1:
                    self._ensure_page_break(word_document)
                continue
            child.render_to_docx(self, word_document, context)

    def _render_title_matter(
        self,
        word_document: WordDocument,
        document: Document,
        context: DocxRenderContext,
    ) -> None:
        self._add_title_line(
            word_document,
            [Text(document.title)],
            font_size=context.theme.title_font_size,
            alignment=context.theme.title_alignment,
            bold=True,
            space_after=12,
        )
        if document.subtitle is not None:
            self._add_title_line(
                word_document,
                document.subtitle,
                font_size=max(context.theme.body_font_size + 1, 12),
                alignment=context.theme.subtitle_alignment,
                italic=True,
                space_after=10,
            )
        for line, is_last_for_author in document.settings.iter_author_title_lines():
            self._add_title_line(
                word_document,
                list(line.fragments),
                font_size=self._title_line_font_size(line, context.theme),
                alignment=self._title_line_alignment(line, context.theme),
                italic=line.kind == "affiliation",
                space_after=10 if is_last_for_author else (4 if line.kind == "name" else 3),
            )

    def _add_title_line(
        self,
        container: object,
        fragments: list[Text],
        *,
        font_size: float,
        alignment: str,
        bold: bool = False,
        italic: bool = False,
        space_after: float = 0,
    ) -> None:
        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[alignment]
        paragraph.paragraph_format.space_after = Pt(space_after)
        base_style = TextStyle(font_size=font_size, bold=bold, italic=italic)
        for fragment in fragments:
            style = base_style.merged(
                TextStyle(
                    font_name=fragment.style.font_name,
                    font_size=fragment.style.font_size,
                    color=fragment.style.color,
                    bold=fragment.style.bold,
                    italic=fragment.style.italic,
                    underline=fragment.style.underline,
                )
            )
            if isinstance(fragment, Hyperlink):
                self._append_hyperlink_runs(
                    paragraph,
                    fragment.target,
                    fragment.label,
                    internal=fragment.internal,
                    style=style,
                    default_size=font_size,
                )
                continue
            run = paragraph.add_run(self._resolve_fragment_text(fragment, None, None))
            self._apply_run_style(run, style, default_size=font_size)

    def _title_line_alignment(self, line: AuthorTitleLine, theme: Theme) -> str:
        if line.kind == "name":
            return theme.author_alignment
        if line.kind == "affiliation":
            return theme.affiliation_alignment
        return theme.author_detail_alignment

    def _title_line_font_size(self, line: AuthorTitleLine, theme: Theme) -> float:
        if line.kind == "name":
            return theme.body_font_size
        if line.kind == "affiliation":
            return max(theme.body_font_size - 0.5, 9)
        return max(theme.body_font_size - 1, 9)

    def _is_paginated_generated_page(self, block: object) -> bool:
        return isinstance(block, (TableList, FigureList, TableOfContents))

    def _should_auto_render_footnotes_page(
        self,
        document: Document,
        render_index: RenderIndex,
    ) -> bool:
        return (
            document.theme.footnote_placement == "document"
            and document.theme.auto_footnotes_page
            and bool(render_index.footnotes)
            and not any(isinstance(child, FootnotesPage) for child in document.body.children)
        )

    def _keep_with_next(self, paragraph: object) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_properties.append(OxmlElement("w:keepNext"))

    def _keep_lines_together(self, paragraph: object) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_properties.append(OxmlElement("w:keepLines"))

    def _prevent_table_row_split(self, table: object) -> None:
        for row in table.rows:
            properties = row._tr.get_or_add_trPr()
            properties.append(OxmlElement("w:cantSplit"))

    def _docx_footnotes_part(self, word_document: WordDocument) -> FootnotesPart:
        if self._native_footnotes_part is not None:
            return self._native_footnotes_part
        try:
            footnotes_part = word_document.part.part_related_by(RT.FOOTNOTES)
        except KeyError:
            footnotes_part = FootnotesPart.default(word_document.part.package)
            word_document.part.relate_to(footnotes_part, RT.FOOTNOTES)
        self._native_footnotes_part = footnotes_part
        return footnotes_part

    def _ensure_native_footnote(
        self,
        fragment: Footnote,
        *,
        theme: Theme,
        render_index: RenderIndex,
        word_document: WordDocument,
    ) -> int:
        footnote_id = render_index.footnote_number(fragment)
        if footnote_id in self._rendered_native_footnotes:
            return footnote_id

        footnotes_part = self._docx_footnotes_part(word_document)
        paragraph = footnotes_part.add_footnote_paragraph(footnote_id)
        paragraph.paragraph_format.space_after = Pt(0)
        reference_run = paragraph.add_run()
        reference_run.font.superscript = True
        reference_run._r.append(OxmlElement("w:footnoteRef"))
        spacer_run = paragraph.add_run(" ")
        self._apply_run_style(
            spacer_run,
            Text("").style,
            default_size=max(theme.body_font_size - 1, 8),
        )
        self._append_runs(
            paragraph,
            fragment.note,
            default_size=max(theme.body_font_size - 1, 8),
            theme=theme,
            render_index=render_index,
            word_document=word_document,
        )
        self._rendered_native_footnotes.add(footnote_id)
        return footnote_id

    def _append_native_footnote_reference(self, run: object, footnote_id: int) -> None:
        reference = OxmlElement("w:footnoteReference")
        reference.set(qn("w:id"), str(footnote_id))
        run._r.append(reference)

    def _ensure_page_break(self, word_document: WordDocument) -> None:
        if self._ends_with_page_break(word_document):
            return
        word_document.add_page_break()

    def _ends_with_page_break(self, word_document: WordDocument) -> bool:
        if not word_document.paragraphs:
            return False
        return 'w:type="page"' in word_document.paragraphs[-1]._p.xml

    def _add_heading(
        self,
        container: object,
        title: list[Text],
        level: int,
        theme: Theme,
        *,
        number_label: str | None = None,
        anchor: str | None = None,
    ) -> None:
        paragraph = self._add_paragraph(container)
        paragraph.style = "Title" if level == 0 else f"Heading {min(level, 9)}"
        if level == 0:
            paragraph.alignment = ALIGNMENTS[theme.title_alignment]
        else:
            paragraph.alignment = ALIGNMENTS[theme.heading_alignment(level)]
            paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            paragraph.paragraph_format.space_after = Pt(10 if level == 1 else 6)
        self._append_runs(
            paragraph,
            self._heading_fragments(title, number_label),
            default_size=theme.title_font_size if level == 0 else theme.heading_size(level),
            theme=theme,
        )
        if anchor is not None:
            self._add_bookmark(paragraph, anchor)

    def _add_paragraph(self, container: object) -> object:
        if self._is_cell_container(container):
            container_id = id(container)
            if container_id not in self._initialized_cells and container.paragraphs:
                self._initialized_cells.add(container_id)
                return container.paragraphs[0]
            self._initialized_cells.add(container_id)
        return container.add_paragraph()

    def _is_cell_container(self, container: object) -> bool:
        return hasattr(container, "_tc") and hasattr(container, "add_table")

    def _assert_document_container(self, container: object, block_name: str) -> None:
        if self._is_cell_container(container):
            raise DocscriptorError(f"{block_name} cannot be rendered inside a Box")

    def _assert_box_child_supported(self, child: object) -> None:
        if isinstance(
            child,
            (
                CommentsPage,
                FootnotesPage,
                ReferencesPage,
                TableOfContents,
                TableList,
                FigureList,
            ),
        ):
            raise DocscriptorError(
                f"{type(child).__name__} cannot be rendered inside a Box"
            )

    def _heading_fragments(self, title: list[Text], number_label: str | None) -> list[Text]:
        if not number_label:
            return title
        return [Text(f"{number_label} ")] + title

    def _apply_paragraph_style(self, paragraph: object, style: ParagraphStyle) -> None:
        paragraph.alignment = ALIGNMENTS[style.alignment]
        if style.space_after is not None:
            paragraph.paragraph_format.space_after = Pt(style.space_after)
        if style.leading is not None:
            paragraph.paragraph_format.line_spacing = Pt(style.leading)

    def _append_runs(
        self,
        paragraph: object,
        fragments: list[Text],
        default_size: float | None = None,
        *,
        theme: Theme | None = None,
        render_index: RenderIndex | None = None,
        word_document: WordDocument | None = None,
    ) -> None:
        for fragment in fragments:
            if isinstance(fragment, Hyperlink):
                self._append_hyperlink_runs(
                    paragraph,
                    fragment.target,
                    fragment.label,
                    internal=fragment.internal,
                    style=fragment.style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, _BlockReference) and theme is not None and render_index is not None:
                anchor = self._block_reference_anchor(fragment.target, render_index)
                self._append_hyperlink_runs(
                    paragraph,
                    anchor,
                    [Text(self._resolve_block_reference(fragment.target, theme, render_index), style=fragment.style)],
                    internal=True,
                    style=fragment.style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, Citation) and render_index is not None:
                self._append_hyperlink_runs(
                    paragraph,
                    render_index.citation_anchor(fragment.target),
                    [Text(f"[{render_index.citation_number(fragment.target)}]", style=fragment.style)],
                    internal=True,
                    style=fragment.style,
                    default_size=default_size,
                )
                continue
            if isinstance(fragment, Comment):
                self._append_comment_runs(
                    paragraph,
                    fragment,
                    default_size=default_size,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                continue
            if isinstance(fragment, Footnote):
                self._append_footnote_runs(
                    paragraph,
                    fragment,
                    default_size=default_size,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )
                continue
            if isinstance(fragment, Math):
                self._append_math_runs(paragraph, fragment, default_size=default_size)
                continue
            run = paragraph.add_run(self._resolve_fragment_text(fragment, theme, render_index))
            self._apply_run_style(run, fragment.style, default_size=default_size)

    def _append_hyperlink_runs(
        self,
        paragraph: object,
        target: str | None,
        label_fragments: list[Text],
        *,
        internal: bool,
        style: TextStyle,
        default_size: float | None,
    ) -> None:
        if not target:
            self._append_runs(
                paragraph,
                label_fragments,
                default_size=default_size,
            )
            return

        hyperlink = OxmlElement("w:hyperlink")
        if internal:
            hyperlink.set(qn("w:anchor"), target)
        else:
            relationship_id = paragraph.part.relate_to(
                target,
                RT.HYPERLINK,
                is_external=True,
            )
            hyperlink.set(qn("r:id"), relationship_id)

        label_text = "".join(fragment.plain_text() for fragment in label_fragments)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        if style.font_name is not None:
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), style.font_name)
            fonts.set(qn("w:hAnsi"), style.font_name)
            run_properties.append(fonts)
        font_size = style.font_size if style.font_size is not None else default_size
        if font_size is not None:
            size = OxmlElement("w:sz")
            size.set(qn("w:val"), str(int(round(font_size * 2))))
            run_properties.append(size)
        if style.bold is not None:
            bold = OxmlElement("w:b")
            if not style.bold:
                bold.set(qn("w:val"), "0")
            run_properties.append(bold)
        if style.italic is not None:
            italic = OxmlElement("w:i")
            if not style.italic:
                italic.set(qn("w:val"), "0")
            run_properties.append(italic)
        if style.color is not None:
            color = OxmlElement("w:color")
            color.set(qn("w:val"), style.color)
            run_properties.append(color)
        if style.underline:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            run_properties.append(underline)
        run.append(run_properties)
        text = OxmlElement("w:t")
        text.text = label_text
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _add_bookmark(self, paragraph: object, anchor: str) -> None:
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(self._bookmark_id))
        bookmark_start.set(qn("w:name"), anchor)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(self._bookmark_id))
        insert_index = (
            1
            if len(paragraph._p) > 0 and paragraph._p[0].tag == qn("w:pPr")
            else 0
        )
        paragraph._p.insert(insert_index, bookmark_start)
        paragraph._p.append(bookmark_end)
        self._bookmark_id += 1

    def _block_reference_anchor(
        self,
        target: Table | Figure,
        render_index: RenderIndex,
    ) -> str | None:
        if isinstance(target, Table):
            return render_index.table_anchor(target)
        return render_index.figure_anchor(target)

    def _apply_run_style(self, run: object, style: object, *, default_size: float | None = None) -> None:
        font = run.font
        if style.font_name:
            font.name = style.font_name
        if style.font_size is not None:
            font.size = Pt(style.font_size)
        elif default_size is not None:
            font.size = Pt(default_size)
        if style.bold is not None:
            font.bold = style.bold
        if style.italic is not None:
            font.italic = style.italic
        if style.underline is not None:
            font.underline = style.underline
        if style.color is not None:
            font.color.rgb = RGBColor.from_string(style.color)

    def _append_comment_runs(
        self,
        paragraph: object,
        fragment: Comment,
        *,
        default_size: float | None,
        theme: Theme | None,
        render_index: RenderIndex | None,
        word_document: WordDocument | None,
    ) -> None:
        visible_runs: list[object] = []
        if fragment.value:
            visible_run = paragraph.add_run(fragment.value)
            self._apply_run_style(visible_run, fragment.style, default_size=default_size)
            visible_runs.append(visible_run)

        marker_run = paragraph.add_run(self._comment_marker(fragment, render_index))
        self._apply_run_style(marker_run, fragment.style, default_size=max((default_size or 10.0) - 2, 8))
        marker_run.font.superscript = True
        anchor_runs = visible_runs or [marker_run]
        if word_document is not None and render_index is not None:
            word_document.add_comment(
                anchor_runs,
                text=self._flatten_fragments(fragment.comment, theme, render_index),
                author=fragment.author or "",
                initials=fragment.initials,
            )

    def _append_footnote_runs(
        self,
        paragraph: object,
        fragment: Footnote,
        *,
        default_size: float | None,
        theme: Theme | None,
        render_index: RenderIndex | None,
        word_document: WordDocument | None,
    ) -> None:
        if fragment.value:
            visible_run = paragraph.add_run(fragment.value)
            self._apply_run_style(visible_run, fragment.style, default_size=default_size)

        if (
            theme is not None
            and word_document is not None
            and render_index is not None
            and theme.footnote_placement == "page"
        ):
            marker_run = paragraph.add_run()
            self._apply_run_style(
                marker_run,
                fragment.style,
                default_size=max((default_size or 10.0) - 2, 8),
            )
            marker_run.font.superscript = True
            self._append_native_footnote_reference(
                marker_run,
                self._ensure_native_footnote(
                    fragment,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                ),
            )
            return

        marker_run = paragraph.add_run(self._footnote_marker(fragment, render_index))
        self._apply_run_style(marker_run, fragment.style, default_size=max((default_size or 10.0) - 2, 8))
        marker_run.font.superscript = True

    def _append_math_runs(self, paragraph: object, fragment: Math, *, default_size: float | None = None) -> None:
        for segment in parse_latex_segments(fragment.value):
            run = paragraph.add_run(segment.text)
            self._apply_run_style(run, fragment.style, default_size=default_size)
            if segment.vertical_align == SUPERSCRIPT:
                run.font.superscript = True
            elif segment.vertical_align == SUBSCRIPT:
                run.font.subscript = True

    def _render_list(
        self,
        container: object,
        list_block: BulletList | NumberedList,
        theme: Theme,
        render_index: RenderIndex,
        *,
        word_document: WordDocument,
    ) -> None:
        list_style = list_block.style or theme.list_style(ordered=isinstance(list_block, NumberedList))
        for index, item in enumerate(list_block.items):
            paragraph = self._add_paragraph(container)
            self._apply_paragraph_style(paragraph, item.style)
            paragraph.paragraph_format.left_indent = Inches(list_style.indent)
            paragraph.paragraph_format.first_line_indent = Inches(-list_style.indent)
            marker = list_style.marker_for(index)
            if marker:
                marker_run = paragraph.add_run(f"{marker} ")
                self._apply_run_style(marker_run, Text("").style, default_size=theme.body_font_size)
            self._append_runs(paragraph, item.content, theme=theme, render_index=render_index, word_document=word_document)

    def _render_code_block(self, container: object, code_block: CodeBlock, theme: Theme) -> None:
        if code_block.language:
            label = self._add_paragraph(container)
            label.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label.paragraph_format.space_after = Pt(2)
            run = label.add_run(code_block.language.upper())
            run.font.name = theme.monospace_font_name
            run.font.size = Pt(theme.caption_size())
            run.font.bold = True

        paragraph = self._add_paragraph(container)
        self._apply_paragraph_style(paragraph, code_block.style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.right_indent = Inches(0.1)
        paragraph.paragraph_format.space_before = Pt(6)
        self._set_paragraph_shading(paragraph, "F5F7FA")

        run = paragraph.add_run()
        run.font.name = theme.monospace_font_name
        run.font.size = Pt(max(theme.body_font_size - 1, 8))

        for index, line in enumerate(code_block.code.replace("\r\n", "\n").replace("\r", "\n").split("\n")):
            if index:
                run.add_break()
            run.add_text(line)

    def _render_equation(self, container: object, equation: Equation, theme: Theme) -> None:
        paragraph = self._add_paragraph(container)
        self._apply_paragraph_style(paragraph, equation.style)
        if paragraph.alignment is None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._append_math_runs(
            paragraph,
            Math(equation.expression),
            default_size=max(theme.body_font_size + 1, 12),
        )

    def _render_box(
        self,
        container: object,
        box: Box,
        theme: Theme,
        render_index: RenderIndex,
        settings: object,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        outer_table = container.add_table(rows=1, cols=1)
        outer_table.alignment = TABLE_ALIGNMENTS[theme.box_alignment]
        cell = outer_table.rows[0].cells[0]
        cell._tc.clear_content()
        self._initialized_cells.discard(id(cell))
        self._set_cell_shading(cell, box.style.background_color)
        self._set_cell_borders(cell, box.style.border_color, box.style.border_width)
        self._set_cell_padding(cell, box.style.padding)

        if box.title is not None:
            title_paragraph = self._add_paragraph(cell)
            title_paragraph.paragraph_format.space_after = Pt(6)
            if box.style.title_background_color is not None:
                self._set_paragraph_shading(title_paragraph, box.style.title_background_color)
            self._append_runs(
                title_paragraph,
                box.title,
                default_size=theme.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

        context = DocxRenderContext(
            theme=theme,
            render_index=render_index,
            settings=settings,
            unit=unit,
            word_document=word_document,
        )
        for child in box.children:
            self._assert_box_child_supported(child)
            self._render_block(cell, child, context)

        if not cell.paragraphs:
            cell.add_paragraph()

    def _render_sheet(
        self,
        container: object,
        sheet: Sheet,
        theme: Theme,
        render_index: RenderIndex,
        settings: object,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        outer_table = container.add_table(rows=1, cols=1)
        outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = outer_table.rows[0].cells[0]
        cell._tc.clear_content()
        self._initialized_cells.discard(id(cell))
        self._set_cell_shading(cell, sheet.background_color)
        if sheet.border_color is not None and sheet.border_width > 0:
            self._set_cell_borders(cell, sheet.border_color, sheet.border_width)
        self._set_cell_padding(cell, 8)

        visible_items = sorted(
            (
                (index, item)
                for index, item in enumerate(sheet.items)
                if isinstance(item, (ImageBox, TextBox))
            ),
            key=lambda indexed: (indexed[1].z_index, indexed[0]),
        )
        for _, item in visible_items:
            if isinstance(item, ImageBox):
                paragraph = self._add_paragraph(cell)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(
                    self._image_box_picture_source(item),
                    width=Inches(length_to_inches(item.width, sheet.unit or unit)),
                    height=Inches(length_to_inches(item.height, sheet.unit or unit)),
                )
            else:
                paragraph = self._add_paragraph(cell)
                paragraph.alignment = ALIGNMENTS[item.align]
                paragraph.paragraph_format.space_after = Pt(2)
                self._append_runs(
                    paragraph,
                    item.content,
                    default_size=item.font_size or theme.body_font_size,
                    theme=theme,
                    render_index=render_index,
                    word_document=word_document,
                )

        if not visible_items:
            cell.add_paragraph()

    def _render_table(
        self,
        container: object,
        table_block: Table,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        def render_caption() -> None:
            if table_block.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.caption_alignment]
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.table_label,
                    render_index.table_number(table_block),
                    table_block.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.table_anchor(table_block)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.table_caption_position == "above":
                self._keep_with_next(caption)

        if table_block.caption is not None and theme.table_caption_position == "above":
            render_caption()

        layout = build_table_layout(table_block.header_rows, table_block.rows)
        table = container.add_table(rows=layout.row_count, cols=layout.column_count)
        table.style = "Table Grid"
        table.alignment = TABLE_ALIGNMENTS[theme.table_alignment]
        self._prevent_table_row_split(table)
        column_widths = table_block.column_widths_in_inches(unit)
        if column_widths is not None:
            for column_index, width in enumerate(column_widths):
                table.columns[column_index].width = Inches(width)

        for placement in layout.placements:
            start_cell = table.cell(placement.row, placement.column)
            target_cell = start_cell
            if placement.cell.colspan > 1 or placement.cell.rowspan > 1:
                end_cell = table.cell(
                    placement.row + placement.cell.rowspan - 1,
                    placement.column + placement.cell.colspan - 1,
                )
                target_cell = start_cell.merge(end_cell)

            paragraph = target_cell.paragraphs[0]
            self._append_runs(
                paragraph,
                placement.cell.content.content or [Text("")],
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            if placement.header:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(table_block.style.header_text_color)
                self._set_cell_shading(target_cell, table_block.style.header_background_color)
            else:
                background_color = placement.cell.background_color
                if background_color is None and table_block.style.alternate_row_background_color is not None and placement.body_row_index is not None and placement.body_row_index % 2 == 1:
                    background_color = table_block.style.alternate_row_background_color
                if background_color is None:
                    background_color = table_block.style.body_background_color
                if background_color is not None:
                    self._set_cell_shading(target_cell, background_color)
            self._set_cell_borders(target_cell, table_block.style.border_color, 0.5)
            self._set_cell_padding(target_cell, table_block.style.cell_padding)

        if table_block.caption is not None and theme.table_caption_position == "below":
            render_caption()

    def _render_figure(
        self,
        container: object,
        figure: Figure,
        theme: Theme,
        render_index: RenderIndex,
        unit: str,
        *,
        word_document: WordDocument,
    ) -> None:
        def render_caption() -> None:
            if figure.caption is None:
                return
            caption = self._add_paragraph(container)
            caption.alignment = ALIGNMENTS[theme.caption_alignment]
            self._keep_lines_together(caption)
            self._append_runs(
                caption,
                self._caption_fragments(
                    theme.figure_label,
                    render_index.figure_number(figure),
                    figure.caption,
                ),
                default_size=theme.caption_size(),
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            anchor = render_index.figure_anchor(figure)
            if anchor is not None:
                self._add_bookmark(caption, anchor)
            if theme.figure_caption_position == "above":
                self._keep_with_next(caption)

        if figure.caption is not None and theme.figure_caption_position == "above":
            render_caption()

        paragraph = self._add_paragraph(container)
        paragraph.alignment = ALIGNMENTS[theme.figure_alignment]
        if figure.caption is not None and theme.figure_caption_position == "below":
            self._keep_with_next(paragraph)
        run = paragraph.add_run()
        resolved_width = figure.width_in_inches(unit)
        resolved_height = figure.height_in_inches(unit)
        width = Inches(resolved_width) if resolved_width is not None else None
        height = Inches(resolved_height) if resolved_height is not None else None
        run.add_picture(self._figure_picture_source(figure), width=width, height=height)

        if figure.caption is not None and theme.figure_caption_position == "below":
            render_caption()

    def _set_paragraph_shading(self, paragraph: object, fill: str) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        paragraph_properties.append(shading)

    def _set_page_background(self, word_document: WordDocument, fill: str) -> None:
        document_element = word_document._element
        existing = document_element.find(qn("w:background"))
        if existing is not None:
            document_element.remove(existing)
        background = OxmlElement("w:background")
        background.set(qn("w:color"), fill)
        document_element.insert(0, background)

    def _set_cell_shading(self, cell: object, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def _set_cell_borders(self, cell: object, color: str, width: float) -> None:
        properties = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        size = str(max(int(round(width * 8)), 0))
        for edge_name in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), size)
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), color)
            borders.append(edge)
        properties.append(borders)

    def _set_cell_padding(self, cell: object, padding: float) -> None:
        properties = cell._tc.get_or_add_tcPr()
        margins = OxmlElement("w:tcMar")
        margin_value = str(max(int(round(padding * 20)), 0))
        for side in ("top", "left", "bottom", "right"):
            element = OxmlElement(f"w:{side}")
            element.set(qn("w:w"), margin_value)
            element.set(qn("w:type"), "dxa")
            margins.append(element)
        properties.append(margins)

    def _configure_named_style(
        self,
        word_document: WordDocument,
        style_name: str,
        *,
        font_name: str,
        font_size: float,
        bold: bool,
        italic: bool,
    ) -> None:
        style = word_document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)

    def _resolve_fragment_text(self, fragment: Text, theme: Theme | None, render_index: RenderIndex | None) -> str:
        if isinstance(fragment, _BlockReference):
            if theme is None or render_index is None:
                return fragment.plain_text()
            return self._resolve_block_reference(fragment.target, theme, render_index)
        if isinstance(fragment, Citation):
            if render_index is None:
                return fragment.plain_text()
            return f"[{render_index.citation_number(fragment.target)}]"
        if isinstance(fragment, Hyperlink):
            return fragment.plain_text()
        if isinstance(fragment, Comment):
            return fragment.value
        if isinstance(fragment, Footnote):
            return fragment.value
        if isinstance(fragment, Math):
            return fragment.plain_text()
        return fragment.value

    def _resolve_block_reference(self, target: Table | Figure, theme: Theme, render_index: RenderIndex) -> str:
        if isinstance(target, Table):
            number = render_index.table_number(target)
            if number is None:
                raise DocscriptorError("Table references require the target table to have a caption and be included in the document")
            return f"{theme.table_label} {number}"

        number = render_index.figure_number(target)
        if number is None:
            raise DocscriptorError("Figure references require the target figure to have a caption and be included in the document")
        return f"{theme.figure_label} {number}"

    def _caption_fragments(self, label: str, number: int | None, caption: Paragraph) -> list[Text]:
        if number is None:
            return caption.content
        return [Text(f"{label} {number}. ")] + caption.content

    def _render_caption_list(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        entries: list[object],
        theme: Theme,
        render_index: RenderIndex,
        default_title: str,
        label: str,
    ) -> None:
        self._add_heading(word_document, title or [Text(default_title)], level=theme.generated_section_level, theme=theme, number_label=None)
        for entry in entries:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            anchor = entry.anchor
            self._append_hyperlink_runs(
                paragraph,
                anchor,
                self._caption_fragments(label, entry.number, entry.block.caption),
                internal=True,
                style=TextStyle(),
                default_size=theme.caption_size(),
            )

    def _render_comments_page(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, title or [Text(theme.comments_title)], level=theme.generated_section_level, theme=theme, number_label=None)
        for entry in render_index.comments:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            self._append_runs(
                paragraph,
                [Text(f"[{entry.number}] ")] + entry.comment.comment,
                default_size=theme.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

    def _figure_picture_source(self, figure: Figure) -> str | BytesIO:
        source = figure.image_source
        if isinstance(source, Path):
            return str(source)
        if hasattr(source, "savefig"):
            buffer = BytesIO()
            save_kwargs: dict[str, object] = {
                "format": figure.format,
            }
            if figure.dpi is not None:
                save_kwargs["dpi"] = figure.dpi
            source.savefig(buffer, **save_kwargs)
            buffer.seek(0)
            return buffer
        raise TypeError(f"Unsupported figure source for DOCX rendering: {type(source)!r}")

    def _image_box_picture_source(self, image_box: ImageBox) -> str | BytesIO:
        source = image_box.image_source
        if isinstance(source, Path):
            return str(source)
        if hasattr(source, "savefig"):
            buffer = BytesIO()
            save_kwargs: dict[str, object] = {"format": image_box.format}
            if image_box.dpi is not None:
                save_kwargs["dpi"] = image_box.dpi
            source.savefig(buffer, **save_kwargs)
            buffer.seek(0)
            return buffer
        raise TypeError(f"Unsupported image source for DOCX sheet rendering: {type(source)!r}")

    def _render_footnotes_page(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, title or [Text(theme.footnotes_title)], level=theme.generated_section_level, theme=theme, number_label=None)
        for entry in render_index.footnotes:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            self._append_runs(
                paragraph,
                [Text(f"[{entry.number}] ")] + entry.footnote.note,
                default_size=theme.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )

    def _render_references_page(
        self,
        word_document: WordDocument,
        title: list[Text] | None,
        theme: Theme,
        render_index: RenderIndex,
    ) -> None:
        word_document.add_page_break()
        self._add_heading(word_document, title or [Text(theme.references_title)], level=theme.generated_section_level, theme=theme, number_label=None)
        for entry in render_index.citations:
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.first_line_indent = Inches(-0.3)
            self._append_runs(
                paragraph,
                [Text(f"[{entry.number}] ")] + entry.source.reference_fragments(),
                default_size=theme.body_font_size,
                theme=theme,
                render_index=render_index,
                word_document=word_document,
            )
            self._add_bookmark(paragraph, entry.anchor)

    def _render_table_of_contents(
        self,
        word_document: WordDocument,
        block: TableOfContents,
        context: DocxRenderContext,
    ) -> None:
        theme = context.theme
        render_index = context.render_index
        self._add_heading(word_document, block.title or [Text(theme.contents_title)], level=theme.generated_section_level, theme=theme, number_label=None)
        for entry in render_index.headings:
            if not block.includes_level(entry.level):
                continue
            toc_style = self._toc_level_style(block, entry.level)
            paragraph = word_document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(toc_style.indent)
            paragraph.paragraph_format.space_before = Pt(toc_style.space_before)
            paragraph.paragraph_format.space_after = Pt(toc_style.space_after)
            if block.show_page_numbers:
                text_width = context.settings.text_width_in_inches()
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(text_width),
                    WD_TAB_ALIGNMENT.RIGHT,
                    WD_TAB_LEADER.DOTS if block.leader == "." else WD_TAB_LEADER.SPACES,
                )
            self._append_hyperlink_runs(
                paragraph,
                entry.anchor,
                self._heading_fragments(entry.title, entry.number),
                internal=True,
                style=TextStyle(
                    bold=toc_style.bold,
                    italic=toc_style.italic,
                ),
                default_size=theme.body_font_size + toc_style.font_size_delta,
            )
            if block.show_page_numbers and entry.anchor is not None:
                paragraph.add_run("\t")
                self._append_pageref_field(paragraph, entry.anchor)

    def _toc_level_style(self, block: TableOfContents, level: int) -> TocLevelStyle:
        defaults = TocLevelStyle(
            indent=0.24 * max(level - 1, 0),
            space_before=12 if level == 1 else (3 if level == 2 else 0),
            space_after=7 if level == 1 else (3 if level == 2 else 2),
            font_size_delta=0.6 if level == 1 else 0,
            bold=True if level == 1 else False,
            italic=False,
        )
        override = block.style_for_level(level)
        return TocLevelStyle(
            indent=defaults.indent if override.indent is None else override.indent,
            space_before=defaults.space_before if override.space_before is None else override.space_before,
            space_after=defaults.space_after if override.space_after is None else override.space_after,
            font_size_delta=defaults.font_size_delta if override.font_size_delta is None else override.font_size_delta,
            bold=defaults.bold if override.bold is None else override.bold,
            italic=defaults.italic if override.italic is None else override.italic,
        )

    def _configure_page_number_sections(
        self,
        word_document: WordDocument,
        theme: Theme,
        *,
        has_front_matter: bool,
        has_main_matter: bool,
    ) -> None:
        sections = list(word_document.sections)
        if not sections:
            return

        if has_front_matter:
            self._set_section_page_number_format(
                sections[0],
                theme.front_matter_page_number_format,
                start=1,
            )
            self._add_page_number_footer(
                sections[0],
                theme,
                front_matter=True,
            )
            if has_main_matter and len(sections) > 1:
                sections[1].footer.is_linked_to_previous = False
                self._set_section_page_number_format(
                    sections[1],
                    theme.main_matter_page_number_format,
                    start=1,
                )
                self._add_page_number_footer(
                    sections[1],
                    theme,
                    front_matter=False,
                )
            return

        self._set_section_page_number_format(
            sections[0],
            theme.main_matter_page_number_format,
            start=1,
        )
        self._add_page_number_footer(
            sections[0],
            theme,
            front_matter=False,
        )

    def _add_page_number_footer(
        self,
        section: object,
        theme: Theme,
        *,
        front_matter: bool,
    ) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.style = "Footer"
        paragraph.alignment = ALIGNMENTS[theme.page_number_alignment]
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        parts = theme.page_number_format.split("{page}")
        for index, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                self._apply_run_style(
                    run,
                    Text(part).style,
                    default_size=theme.page_number_font_size,
                )
            if index < len(parts) - 1:
                self._append_page_number_field(paragraph)

    def _set_section_page_number_format(
        self,
        section: object,
        page_number_format: str,
        *,
        start: int = 1,
    ) -> None:
        format_map = {
            "decimal": "decimal",
            "lower-roman": "lowerRoman",
            "upper-roman": "upperRoman",
            "lower-alpha": "lowerLetter",
            "upper-alpha": "upperLetter",
        }
        sect_pr = section._sectPr
        page_number_type = sect_pr.find(qn("w:pgNumType"))
        if page_number_type is None:
            page_number_type = OxmlElement("w:pgNumType")
            sect_pr.append(page_number_type)
        page_number_type.set(
            qn("w:fmt"),
            format_map.get(page_number_format, "decimal"),
        )
        page_number_type.set(qn("w:start"), str(start))

    def _append_page_number_field(self, paragraph: object) -> None:
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        field.append(run)
        paragraph._p.append(field)

    def _append_pageref_field(self, paragraph: object, anchor: str) -> None:
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), f"PAGEREF {anchor} \\h")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        field.append(run)
        paragraph._p.append(field)

    def _comment_marker(self, fragment: Comment, render_index: RenderIndex | None) -> str:
        if render_index is None:
            return "[?]"
        return f"[{render_index.comment_number(fragment)}]"

    def _footnote_marker(self, fragment: Footnote, render_index: RenderIndex | None) -> str:
        if render_index is None:
            return "?"
        return str(render_index.footnote_number(fragment))

    def _flatten_fragments(self, fragments: list[Text], theme: Theme | None, render_index: RenderIndex | None) -> str:
        parts: list[str] = []
        for fragment in fragments:
            if isinstance(fragment, Comment):
                parts.append(fragment.value)
                parts.append(self._comment_marker(fragment, render_index))
                continue
            if isinstance(fragment, Footnote):
                parts.append(fragment.value)
                parts.append(self._footnote_marker(fragment, render_index))
                continue
            parts.append(self._resolve_fragment_text(fragment, theme, render_index))
        return "".join(parts)
